"""
Infrastructure: Language Detector — phát hiện ngôn ngữ trong file .docx.
Sử dụng langdetect để xác định nội dung tiếng Nga tự động.
"""
import logging
from typing import Optional
# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from langdetect import detect, LangDetectException

logger = logging.getLogger(__name__)

# Số ký tự tối thiểu cần để detect chính xác
MIN_CHARS_FOR_DETECTION = 50
# Số paragraph tối đa cần đọc
MAX_PARAGRAPHS_TO_READ = 20


def detect_language(docx_path: str) -> Optional[str]:
    """
    Đọc file .docx và phát hiện ngôn ngữ chính.
    Trả về language code (vd: 'ru', 'en', 'vi') hoặc None nếu không detect được.
    """
    try:
        doc = Document(docx_path)
        text_parts = []
        total_chars = 0

        # Đọc text từ paragraphs
        for para in doc.paragraphs[:MAX_PARAGRAPHS_TO_READ]:
            text = para.text.strip()
            if text:
                text_parts.append(text)
                total_chars += len(text)
                if total_chars >= 500:
                    break

        # Nếu paragraphs không đủ, đọc thêm từ tables
        if total_chars < MIN_CHARS_FOR_DETECTION:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text.strip()
                            if text:
                                text_parts.append(text)
                                total_chars += len(text)
                                if total_chars >= 500:
                                    break
                        if total_chars >= 500:
                            break
                    if total_chars >= 500:
                        break
                if total_chars >= 500:
                    break

        if total_chars < MIN_CHARS_FOR_DETECTION:
            logger.warning(
                f"⚠️ File {docx_path} chỉ có {total_chars} ký tự, "
                f"không đủ để phát hiện ngôn ngữ (cần ≥ {MIN_CHARS_FOR_DETECTION})"
            )
            return None

        combined_text = " ".join(text_parts)
        lang = detect(combined_text)
        logger.info(f"🔍 Phát hiện ngôn ngữ: {lang} ({total_chars} ký tự) — {docx_path}")
        return lang

    except LangDetectException as e:
        logger.warning(f"⚠️ Không thể phát hiện ngôn ngữ file {docx_path}: {e}")
        return None
    except Exception as e:
        logger.error(f"❌ Lỗi đọc file để detect ngôn ngữ {docx_path}: {e}")
        return None


def is_russian_document(docx_path: str) -> bool:
    """Kiểm tra xem file .docx có phải tiếng Nga không."""
    lang = detect_language(docx_path)
    if lang == "ru":
        return True
    if lang is None:
        # Nếu không detect được, giả sử là tiếng Nga (cho an toàn, tránh bỏ sót)
        logger.info(f"ℹ️ Không detect được ngôn ngữ, mặc định xử lý: {docx_path}")
        return True
    logger.info(f"⏩ File không phải tiếng Nga (detected: {lang}), bỏ qua: {docx_path}")
    return False
