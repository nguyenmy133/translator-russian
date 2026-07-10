"""
Infrastructure: python-docx Adapter — implements IDocumentParser
Dịch file Word giữ nguyên format (bold, italic, font size, color, tables)
Hỗ trợ dịch theo lô (batching) tối ưu hóa tốc độ và API request.
"""
import logging
import html
import xml.etree.ElementTree as ET
from docx import Document
from app.application.ports.document_port import IDocumentParser, DocumentStats
from typing import Callable, Optional

logger = logging.getLogger(__name__)


class DocxParser(IDocumentParser):
    """
    Đọc .docx, dịch toàn bộ paragraph và cell trong table dưới dạng gộp nhóm (batching),
    giúp tối ưu hóa tần suất gọi API (Rate limit) và tăng tốc độ xử lý.
    """

    def translate_and_save(
        self,
        input_path: str,
        output_path: str,
        translate_fn: Callable[[str], str],
    ) -> DocumentStats:
        doc = Document(input_path)
        
        # 1. Thu thập tất cả các paragraphs cần dịch (bao gồm cả trong table)
        units = []
        
        # Paragraphs chính của văn bản
        for para in doc.paragraphs:
            original = para.text.strip()
            if original:
                units.append(self._build_translation_unit(para))
                
        # Paragraphs trong các bảng biểu
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original = para.text.strip()
                        if original:
                            units.append(self._build_translation_unit(para))

        total_paragraphs = len(units)
        total_chars = sum(len(u["original_text"]) for u in units)

        if not units:
            doc.save(output_path)
            return DocumentStats(paragraph_count=0, char_count=0)

        # 2. Thực hiện dịch thuật (Ưu tiên dịch theo lô nếu adapter hỗ trợ)
        translator_adapter = getattr(translate_fn, "__self__", None)
        if translator_adapter and hasattr(translator_adapter, "translate_batch_xml"):
            logger.info(f"⚡ Đang dịch theo lô (Batching 30 đoạn/request) cho {total_paragraphs} đoạn...")
            self._translate_units_in_batches(units, translator_adapter, translate_fn, batch_size=30)
        else:
            logger.info(f"ℹ️ Dịch từng đoạn độc lập (không hỗ trợ batching) cho {total_paragraphs} đoạn...")
            for unit in units:
                self._translate_unit_single(unit, translate_fn)

        # 3. Ghi kết quả dịch ngược lại vào tài liệu Word
        for unit in units:
            self._write_translated_unit(unit, translate_fn)

        doc.save(output_path)
        logger.info(
            f"✅ Đã lưu file dịch: {output_path} "
            f"({total_paragraphs} đoạn, {total_chars:,} ký tự)"
        )
        return DocumentStats(
            paragraph_count=total_paragraphs,
            char_count=total_chars,
        )

    def _build_translation_unit(self, para) -> dict:
        """Trích xuất thông tin định dạng và tạo đơn vị dịch thuật."""
        original = para.text.strip()
        non_empty_runs = [r for r in para.runs if r.text]
        
        unit = {
            "para": para,
            "original_text": original,
            "original_runs_data": [],
            "html_string": "",
            "type": "plain",
            "translated_content": None
        }
        
        if not non_empty_runs:
            unit["type"] = "plain"
            return unit
            
        # Trích xuất dữ liệu định dạng của các runs gốc
        for run in non_empty_runs:
            font_color = None
            try:
                if run.font.color and run.font.color.type:
                    font_color = run.font.color.rgb
            except Exception:
                pass

            style = {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size,
                "font_color": font_color,
            }
            unit["original_runs_data"].append(style)
            
        if len(non_empty_runs) == 1:
            unit["type"] = "single_run"
        else:
            unit["type"] = "multi_run"
            # Tạo chuỗi HTML span mapping
            html_parts = []
            for idx, run in enumerate(non_empty_runs):
                escaped_text = html.escape(run.text)
                html_parts.append(f'<span id="{idx}">{escaped_text}</span>')
            unit["html_string"] = "".join(html_parts)
            
        return unit

    def _translate_units_in_batches(self, units: list, translator_adapter, translate_fn, batch_size: int = 30):
        """Dịch danh sách đơn vị dịch theo từng lô để tránh rate limit."""
        for i in range(0, len(units), batch_size):
            chunk = units[i:i + batch_size]
            
            # Xây dựng XML payload cho cả lô
            xml_parts = ["<doc>"]
            for idx, unit in enumerate(chunk):
                content = unit["html_string"] if unit["type"] == "multi_run" else unit["original_text"]
                if unit["type"] != "multi_run":
                    content = html.escape(content)
                xml_parts.append(f'<u i="{idx}">{content}</u>')
            xml_parts.append("</doc>")
            xml_payload = "".join(xml_parts)
            
            try:
                # Gọi API dịch theo lô
                translated_xml = translator_adapter.translate_batch_xml(xml_payload)
                translated_map = self._parse_translated_xml(translated_xml)
                
                # Gán kết quả dịch
                for idx, unit in enumerate(chunk):
                    val = translated_map.get(str(idx))
                    if val is not None:
                        unit["translated_content"] = val
                    else:
                        # Fallback dịch riêng lẻ nếu bị thiếu thẻ
                        logger.warning(f"⚠️ Thiếu kết quả dịch của đoạn #{i + idx}. Tiến hành dịch riêng lẻ...")
                        self._translate_unit_single(unit, translate_fn)
            except Exception as batch_err:
                logger.error(f"❌ Lỗi dịch lô từ đoạn {i} đến {i+len(chunk)}: {batch_err}. Dịch từng đoạn đơn...")
                for unit in chunk:
                    self._translate_unit_single(unit, translate_fn)

    def _translate_unit_single(self, unit: dict, translate_fn: Callable[[str], str]):
        """Dịch riêng lẻ một đơn vị dịch."""
        try:
            content = unit["html_string"] if unit["type"] == "multi_run" else unit["original_text"]
            unit["translated_content"] = translate_fn(content)
        except Exception as e:
            logger.error(f"❌ Lỗi dịch riêng lẻ đoạn văn: {e}")
            unit["translated_content"] = unit["original_text"]  # Giữ nguyên bản gốc nếu lỗi

    def _parse_translated_xml(self, translated_xml: str) -> dict[str, str]:
        """Phân tích XML trả về từ Gemini để bóc tách bản dịch của từng đoạn."""
        xml_str = translated_xml.strip()
        if not xml_str.startswith("<doc>"):
            xml_str = f"<doc>{xml_str}</doc>"
            
        try:
            root = ET.fromstring(xml_str)
            res = {}
            for child in root:
                if child.tag == "u" and "i" in child.attrib:
                    res[child.attrib["i"]] = self._get_inner_html(child)
            return res
        except Exception as parse_err:
            logger.error(f"⚠️ Lỗi phân tích cú pháp XML trả về: {parse_err}")
            raise parse_err

    def _get_inner_html(self, child) -> str:
        """Lấy nội dung bên trong thẻ XML (bao gồm cả các thẻ con như span)."""
        raw_str = ET.tostring(child, encoding='unicode')
        start_idx = raw_str.find('>') + 1
        end_idx = raw_str.rfind('<')
        if start_idx > 0 and end_idx > start_idx:
            return raw_str[start_idx:end_idx]
        return child.text or ""

    def _write_translated_unit(self, unit: dict, translate_fn: Callable[[str], str]):
        """Ghi nội dung đã dịch đè lên paragraph của tài liệu Word."""
        para = unit["para"]
        translated_content = unit["translated_content"]
        original_runs_data = unit["original_runs_data"]
        
        # Nếu chưa được dịch thành công, giữ nguyên
        if translated_content is None:
            return
            
        if unit["type"] == "plain":
            para.text = translated_content
        elif unit["type"] == "single_run":
            for r in para.runs:
                r.text = ""
            new_run = para.add_run(translated_content)
            self._apply_style(new_run, original_runs_data[0])
        elif unit["type"] == "multi_run":
            try:
                # Tạo XML giả lập để parse thẻ span
                xml_str = f"<div>{translated_content}</div>"
                root = ET.fromstring(xml_str)
                
                # Xoá nội dung cũ
                for r in para.runs:
                    r.text = ""
                    
                # Ghi text trước thẻ span đầu tiên
                if root.text:
                    new_run = para.add_run(root.text)
                    self._apply_style(new_run, original_runs_data[0])
                    
                # Duyệt các span và tail
                for child in root:
                    if child.tag == "span" and "id" in child.attrib:
                        try:
                            run_idx = int(child.attrib["id"])
                            run_style = original_runs_data[run_idx]
                        except (ValueError, IndexError):
                            run_style = original_runs_data[0]
                            
                        # Nội dung text trong span
                        new_run = para.add_run(child.text or "")
                        self._apply_style(new_run, run_style)
                        
                        # Nội dung text đứng sau span
                        if child.tail:
                            new_run_tail = para.add_run(child.tail)
                            self._apply_style(new_run_tail, run_style)
            except Exception as e:
                logger.warning(f"⚠️ Định dạng HTML của đoạn multi-run bị sai lệch: {e}. Fallback dịch thô...")
                try:
                    # Nếu lỗi parse HTML (do LLM làm biến dạng tag), tiến hành dịch thô không định dạng
                    translated_plain = translate_fn(unit["original_text"])
                    for r in para.runs:
                        r.text = ""
                    new_run = para.add_run(translated_plain)
                    self._apply_style(new_run, original_runs_data[0])
                except Exception as fallback_err:
                    logger.error(f"❌ Fallback dịch thô đoạn multi-run thất bại: {fallback_err}")

    @staticmethod
    def _apply_style(run, style_data: dict) -> None:
        run.bold = style_data["bold"]
        run.italic = style_data["italic"]
        run.underline = style_data["underline"]
        if style_data["font_name"]:
            run.font.name = style_data["font_name"]
        if style_data["font_size"]:
            run.font.size = style_data["font_size"]
        if style_data["font_color"]:
            try:
                run.font.color.rgb = style_data["font_color"]
            except Exception:
                pass
